# core/test_runner.py
import threading
import tkinter as tk
from tkinter import messagebox
from ui.control_panel import ControlPanel
from core.screenshot import Screenshot
import time
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class TestRunner:
    def __init__(self, logger, excel_handler, root):
        self.logger = logger
        self.excel_handler = excel_handler
        self.root = root

        self.control_panel = None
        self.current_index = 0
        self.cases_df = None
        self.thread = None
        self.stop_event = threading.Event()

    def run(self):
        file_path = self.excel_handler.select_excel_file()
        if not file_path:
            self.logger.log("未选择Excel文件，终止执行")
            return

        if not self.excel_handler.load_excel():
            self.logger.log("Excel加载失败，终止执行")
            return

        self.cases_df = self.excel_handler.get_valid_cases()
        if self.cases_df.empty:
            messagebox.showinfo("提示", "无有效未执行用例")
            self.logger.log("无有效未执行用例，结束流程")
            return

        self.root.withdraw()
        self.logger.log(f"开始执行，共{len(self.cases_df)}条用例")

        self.stop_event.clear()
        self.thread = threading.Thread(target=self._process_cases)
        self.thread.daemon = True
        self.thread.start()

    def _process_cases(self):
        total = len(self.cases_df)
        while self.current_index < total and not self.stop_event.is_set():
            case = self.cases_df.iloc[self.current_index]
            case_name = str(case.iloc[0])
            verify_point = str(case.iloc[1])

            self.logger.log(f"执行用例[{self.current_index+1}/{total}]: {case_name} - {verify_point}")

            self._show_control_panel(case_name, verify_point)

            self.control_panel.wait_event.wait()

            if self.control_panel.exit_flag:
                self.logger.log("用户请求退出执行")
                break

            if self.control_panel.skip_flag:
                self.logger.log(f"跳过用例：{case_name}")
                self.current_index += 1
                continue

            if self.control_panel.screenshot_done:
                saved_img_path = self.control_panel.screenshot_path
                self.logger.log(f"截图标注完成，文件：{saved_img_path}")

                # 插入Word逻辑
                word_file = self._add_image_to_word(case_name, verify_point, saved_img_path)
                self.logger.log(f"图片插入Word文档：{word_file}")

                # 更新Excel状态
                self.excel_handler.update_case_status(self.cases_df.index[self.current_index], "已执行")
                self.logger.log(f"用例标记为已执行：{case_name}")

                self.current_index += 1

            self.root.after(0, self.control_panel.destroy)
            self.control_panel = None

        self.logger.log("所有用例执行完毕或已停止")
        self.root.after(0, self._finish_execution)

    def _show_control_panel(self, case_name, verify_point):
        event = threading.Event()

        def create_panel():
            self.control_panel = ControlPanel(self.root, case_name, verify_point, event)

        self.root.after(0, create_panel)
        event.wait()

    def _finish_execution(self):
        messagebox.showinfo("完成", "所有用例已执行完毕")
        self.root.deiconify()
        self.current_index = 0
        self.cases_df = None

    def stop(self):
        self.stop_event.set()
        if self.control_panel:
            self.control_panel.exit_flag = True
            self.control_panel.wait_event.set()

    def _add_image_to_word(self, case_name, verify_point, image_path):
        folder = os.path.join(os.getcwd(), "word_output", case_name)
        os.makedirs(folder, exist_ok=True)
        word_path = os.path.join(folder, f"{case_name}.docx")

        if os.path.exists(word_path):
            doc = Document(word_path)
        else:
            doc = Document()
            doc.add_heading(f"用例名：{case_name}", level=1)
            doc.add_paragraph(f"验证点：{verify_point}")

        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(image_path, width=Inches(6))

        caption = f"截图 - {os.path.basename(image_path)}"
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()
        doc.save(word_path)
        return word_path