# main.py
import os
import sys
import time
import threading
import tkinter as tk
from tkinter import messagebox, ttk
import pandas as pd
from openpyxl import load_workbook
import pyautogui
import keyboard
from PIL import Image, ImageTk, ImageDraw, ImageFont
import docx
from docx.shared import Inches
import logging
from datetime import datetime
import shutil

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('automation_test.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)

class AutomationTestTool:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("桌面自动化测试工具")
        self.root.geometry("400x200")
        self.root.resizable(False, False)
        
        # 初始化变量
        self.excel_path = None
        self.df = None
        self.current_case_index = 0
        self.current_case_data = None
        self.screenshot_path = None
        self.is_running = False
        self.skip_current = False
        
        # 创建控制面板
        self.control_panel = None
        self.log_text = None
        self.screenshot_label = None
        self.screenshot_image = None
        
        # 创建主界面
        self.create_main_ui()
        
        # 注册快捷键
        keyboard.on_press_key("f8", self.on_screenshot_hotkey)
        
        # 查找最新的Excel文件
        self.find_latest_excel()
        
    def create_main_ui(self):
        """创建主界面"""
        # 主界面按钮
        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=20)
        
        self.start_button = tk.Button(
            button_frame, 
            text="开始执行", 
            command=self.start_execution,
            width=10,
            height=2
        )
        self.start_button.pack(side=tk.LEFT, padx=10)
        
        self.screenshot_button = tk.Button(
            button_frame, 
            text="截图", 
            command=self.trigger_screenshot,
            width=10,
            height=2,
            state=tk.DISABLED
        )
        self.screenshot_button.pack(side=tk.LEFT, padx=10)
        
        self.exit_button = tk.Button(
            button_frame, 
            text="退出", 
            command=self.exit_app,
            width=10,
            height=2
        )
        self.exit_button.pack(side=tk.LEFT, padx=10)
        
        # 进度标签
        self.progress_label = tk.Label(
            self.root, 
            text="准备就绪", 
            font=("Arial", 12)
        )
        self.progress_label.pack(pady=10)
        
    def find_latest_excel(self):
        """查找最新的Excel文件"""
        excel_dir = "excel_input"
        if not os.path.exists(excel_dir):
            os.makedirs(excel_dir)
            messagebox.showinfo("提示", f"请在 {excel_dir} 目录下放入Excel文件")
            return
            
        excel_files = [f for f in os.listdir(excel_dir) if f.endswith(('.xlsx', '.xls'))]
        if not excel_files:
            messagebox.showinfo("提示", f"{excel_dir} 目录下没有找到Excel文件")
            return
            
        # 按修改时间排序，获取最新的文件
        latest_file = max(excel_files, key=lambda f: os.path.getmtime(os.path.join(excel_dir, f)))
        self.excel_path = os.path.join(excel_dir, latest_file)
        self.progress_label.config(text=f"找到文件: {latest_file}")
        
        # 读取Excel数据
        try:
            self.df = pd.read_excel(self.excel_path)
            # 确保必要的列存在
            required_columns = ['用例文件名', '验证点描述', '执行状态']
            for col in required_columns:
                if col not in self.df.columns:
                    self.df[col] = ''
            # 填充空值
            self.df.fillna('', inplace=True)
        except Exception as e:
            messagebox.showerror("错误", f"读取Excel文件失败: {str(e)}")
            logging.error(f"读取Excel文件失败: {str(e)}")
    
    def start_execution(self):
        """开始执行测试用例"""
        if self.df is None or self.excel_path is None:
            messagebox.showwarning("警告", "未找到有效的Excel文件")
            return
            
        # 隐藏主界面
        self.root.withdraw()
        
        # 创建控制面板
        self.create_control_panel()
        
        # 启动执行线程
        self.is_running = True
        execution_thread = threading.Thread(target=self.execution_loop, daemon=True)
        execution_thread.start()
        
    def execution_loop(self):
        """执行测试用例的主循环"""
        try:
            # 查找未执行的用例
            unexecuted_cases = self.df[self.df['执行状态'] != '已执行'].index.tolist()
            
            if not unexecuted_cases:
                self.log_message("没有找到未执行的用例")
                self.root.after(0, lambda: messagebox.showinfo("提示", "所有用例均已执行完成"))
                return
                
            self.log_message(f"找到 {len(unexecuted_cases)} 个未执行用例")
            
            for idx in unexecuted_cases:
                if not self.is_running:
                    break
                    
                self.current_case_index = idx
                self.current_case_data = self.df.iloc[idx]
                
                case_name = self.current_case_data['用例文件名']
                validation_point = self.current_case_data['验证点描述']
                
                # 更新进度
                progress_text = f"正在执行: {case_name} ({unexecuted_cases.index(idx)+1}/{len(unexecuted_cases)})"
                self.root.after(0, lambda t=progress_text: self.update_progress(t))
                self.log_message(f"开始执行用例: {case_name}")
                self.log_message(f"验证点: {validation_point}")
                
                # 创建用例文件夹
                case_folder = os.path.join("screenshots", case_name)
                if not os.path.exists(case_folder):
                    os.makedirs(case_folder)
                
                # 显示验证点弹窗
                self.root.after(0, lambda n=case_name, v=validation_point: self.show_validation_dialog(n, v))
                
                # 等待截图
                self.wait_for_screenshot()
                
                if self.skip_current:
                    self.skip_current = False
                    self.log_message(f"跳过用例: {case_name}")
                    continue
                
                # 更新Excel状态
                self.update_excel_status(idx, "已执行")
                
                # 询问是否继续当前用例
                continue_current = self.ask_continue_current_case()
                if not continue_current:
                    self.log_message(f"完成用例: {case_name}")
                    
        except Exception as e:
            error_msg = f"执行过程中发生错误: {str(e)}"
            self.log_message(error_msg)
            logging.error(error_msg, exc_info=True)
            self.root.after(0, lambda: messagebox.showerror("错误", error_msg))
        finally:
            self.is_running = False
            self.root.after(0, self.on_execution_finished)
    
    def show_validation_dialog(self, case_name, validation_point):
        """显示验证点弹窗"""
        dialog = tk.Toplevel(self.root)
        dialog.title("验证点信息")
        dialog.geometry("400x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        tk.Label(dialog, text=f"用例名称: {case_name}", font=("Arial", 12, "bold")).pack(pady=10)
        tk.Label(dialog, text=f"验证点: {validation_point}", wraplength=380, justify=tk.LEFT).pack(pady=10)
        
        btn_frame = tk.Frame(dialog)
        btn_frame.pack(pady=20)
        
        tk.Button(btn_frame, text="确定", command=dialog.destroy).pack(side=tk.LEFT, padx=10)
        
        # 居中显示
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (dialog.winfo_width() // 2)
        y = (dialog.winfo_screenheight() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")
        
        dialog.focus_set()
    
    def wait_for_screenshot(self):
        """等待用户触发截图"""
        self.screenshot_event = threading.Event()
        self.root.after(0, lambda: self.screenshot_button.config(state=tk.NORMAL))
        self.screenshot_event.wait()
        self.root.after(0, lambda: self.screenshot_button.config(state=tk.DISABLED))
    
    def trigger_screenshot(self):
        """触发截图流程"""
        if not self.is_running:
            return
            
        # 在后台线程中执行截图
        screenshot_thread = threading.Thread(target=self.take_screenshot, daemon=True)
        screenshot_thread.start()
    
    def on_screenshot_hotkey(self, event):
        """截图快捷键回调"""
        if event.name == "f8" and self.is_running:
            self.trigger_screenshot()
    
    def take_screenshot(self):
        """执行截图和标注流程"""
        try:
            self.log_message("开始截图...")
            
            # 全屏截图
            screenshot = pyautogui.screenshot()
            
            # 保存临时截图
            temp_screenshot_path = f"temp_screenshot_{int(time.time())}.png"
            screenshot.save(temp_screenshot_path)
            
            # 显示标注窗口
            self.root.after(0, lambda p=temp_screenshot_path: self.show_annotation_window(p))
            
        except Exception as e:
            error_msg = f"截图失败: {str(e)}"
            self.log_message(error_msg)
            logging.error(error_msg, exc_info=True)
    
    def show_annotation_window(self, image_path):
        """显示标注窗口"""
        try:
            # 创建标注窗口
            annotation_window = AnnotationWindow(self.root, image_path, self.on_annotation_complete)
            annotation_window.show()
        except Exception as e:
            error_msg = f"显示标注窗口失败: {str(e)}"
            self.log_message(error_msg)
            logging.error(error_msg, exc_info=True)
    
    def on_annotation_complete(self, annotated_image_path):
        """标注完成回调"""
        try:
            case_name = self.current_case_data['用例文件名']
            validation_point = self.current_case_data['验证点描述']
            
            # 保存截图到用例文件夹
            case_folder = os.path.join("screenshots", case_name)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            final_screenshot_path = os.path.join(case_folder, f"screenshot_{timestamp}.png")
            shutil.copy(annotated_image_path, final_screenshot_path)
            
            # 更新日志和预览
            self.log_message(f"截图已保存: {final_screenshot_path}")
            self.root.after(0, lambda p=final_screenshot_path: self.update_screenshot_preview(p))
            
            # 插入到Word文档
            self.insert_to_word(case_name, validation_point, final_screenshot_path)
            
            # 设置事件，继续执行
            self.screenshot_event.set()
            
            # 删除临时文件
            if os.path.exists(annotated_image_path):
                os.remove(annotated_image_path)
                
        except Exception as e:
            error_msg = f"处理标注图像失败: {str(e)}"
            self.log_message(error_msg)
            logging.error(error_msg, exc_info=True)
    
    def insert_to_word(self, case_name, validation_point, screenshot_path):
        """插入截图到Word文档"""
        try:
            # 创建Word文档路径
            word_dir = "test_reports"
            if not os.path.exists(word_dir):
                os.makedirs(word_dir)
                
            word_path = os.path.join(word_dir, f"{case_name}.docx")
            
            # 如果文档不存在则创建新文档
            if not os.path.exists(word_path):
                doc = docx.Document()
                doc.add_heading(f'测试用例: {case_name}', 0)
            else:
                doc = docx.Document(word_path)
            
            # 添加验证点和截图
            doc.add_heading(f'验证点: {validation_point}', level=1)
            doc.add_paragraph(f'截图时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_picture(screenshot_path, width=Inches(6))
            doc.add_page_break()
            
            # 保存文档
            doc.save(word_path)
            self.log_message(f"Word文档已更新: {word_path}")
            
        except Exception as e:
            error_msg = f"更新Word文档失败: {str(e)}"
            self.log_message(error_msg)
            logging.error(error_msg, exc_info=True)
    
    def update_excel_status(self, index, status):
        """更新Excel中的执行状态"""
        try:
            # 使用openpyxl更新特定单元格
            wb = load_workbook(self.excel_path)
            ws = wb.active
            
            # 查找"执行状态"列
            status_col = None
            for col in range(1, ws.max_column + 1):
                if ws.cell(row=1, column=col).value == '执行状态':
                    status_col = col
                    break
            
            if status_col is None:
                # 如果没有执行状态列，添加一列
                status_col = ws.max_column + 1
                ws.cell(row=1, column=status_col).value = '执行状态'
            
            # 更新状态
            ws.cell(row=index + 2, column=status_col).value = status  # +2因为Excel行从1开始且有标题行
            
            # 保存文件
            wb.save(self.excel_path)
            
            # 更新内存中的DataFrame
            self.df.at[index, '执行状态'] = status
            
            self.log_message(f"Excel状态已更新为: {status}")
            
        except Exception as e:
            error_msg = f"更新Excel状态失败: {str(e)}"
            self.log_message(error_msg)
            logging.error(error_msg, exc_info=True)
    
    def ask_continue_current_case(self):
        """询问是否继续当前用例"""
        result = [False]  # 使用列表以便在内部函数中修改
        
        def ask():
            answer = messagebox.askyesno(
                "继续确认", 
                f"用例 {self.current_case_data['用例文件名']} 已完成截图\n是否继续当前用例的其他截图？"
            )
            result[0] = answer
        
        self.root.after(0, ask)
        
        # 等待用户响应
        time.sleep(0.1)
        while result[0] is False and 'ask' in locals():
            self.root.update()
            time.sleep(0.1)
            
        return result[0]
    
    def create_control_panel(self):
        """创建控制面板"""
        if self.control_panel is not None:
            return
            
        self.control_panel = tk.Toplevel(self.root)
        self.control_panel.title("控制面板")
        self.control_panel.geometry("600x500")
        self.control_panel.attributes('-topmost', True)  # 置顶显示
        
        # 日志显示区域
        log_frame = tk.Frame(self.control_panel)
        log_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        tk.Label(log_frame, text="执行日志:", font=("Arial", 10, "bold")).pack(anchor=tk.W)
        
        # 创建文本框和滚动条
        text_frame = tk.Frame(log_frame)
        text_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.log_text = tk.Text(
            text_frame, 
            bg="black", 
            fg="green", 
            font=("Consolas", 10),
            wrap=tk.WORD
        )
        scrollbar = tk.Scrollbar(text_frame, orient=tk.VERTICAL, command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=scrollbar.set)
        
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 截图预览区域
        preview_frame = tk.Frame(self.control_panel)
        preview_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Label(preview_frame, text="截图预览:", font=("Arial", 10, "bold")).pack(anchor=tk.W)
        
        self.screenshot_label = tk.Label(preview_frame, text="暂无截图", bg="lightgray", height=8)
        self.screenshot_label.pack(fill=tk.X, pady=5)
        
        # 控制按钮
        button_frame = tk.Frame(self.control_panel)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Button(
            button_frame, 
            text="截图 (F8)", 
            command=self.trigger_screenshot,
            bg="lightblue"
        ).pack(side=tk.LEFT, padx=5)
        
        tk.Button(
            button_frame, 
            text="跳过当前用例", 
            command=self.skip_current_case,
            bg="lightyellow"
        ).pack(side=tk.LEFT, padx=5)
        
        tk.Button(
            button_frame, 
            text="停止执行", 
            command=self.stop_execution,
            bg="lightcoral"
        ).pack(side=tk.LEFT, padx=5)
        
    def log_message(self, message):
        """记录日志消息"""
        if self.log_text is None:
            return
            
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {message}\n"
        
        # 在主线程中更新UI
        self.root.after(0, lambda: self.append_log_message(log_entry))
        
        # 写入日志文件
        logging.info(message)
    
    def append_log_message(self, message):
        """在日志文本框中添加消息"""
        if self.log_text is None:
            return
            
        self.log_text.insert(tk.END, message)
        self.log_text.see(tk.END)
    
    def update_progress(self, text):
        """更新进度标签"""
        if self.control_panel and self.control_panel.winfo_exists():
            self.control_panel.title(f"控制面板 - {text}")
        self.log_message(text)
    
    def update_screenshot_preview(self, image_path):
        """更新截图预览"""
        if self.screenshot_label is None:
            return
            
        try:
            # 加载并缩放图像
            image = Image.open(image_path)
            image.thumbnail((500, 100))  # 缩放到合适大小
            self.screenshot_image = ImageTk.PhotoImage(image)
            self.screenshot_label.config(image=self.screenshot_image, text="")
        except Exception as e:
            self.screenshot_label.config(text=f"预览加载失败: {str(e)}", image=None)
    
    def skip_current_case(self):
        """跳过当前用例"""
        if self.is_running:
            self.skip_current = True
            self.screenshot_event.set()
            self.log_message(f"跳过用例: {self.current_case_data['用例文件名']}")
    
    def stop_execution(self):
        """停止执行"""
        if messagebox.askyesno("确认", "确定要停止执行吗？"):
            self.is_running = False
            self.screenshot_event.set()  # 确保不会卡在等待截图
            self.log_message("用户停止执行")
    
    def on_execution_finished(self):
        """执行完成后的处理"""
        self.log_message("所有用例执行完成")
        if self.control_panel:
            self.control_panel.destroy()
            self.control_panel = None
            
        self.root.deiconify()  # 显示主界面
        messagebox.showinfo("完成", "所有用例执行完成")
    
    def exit_app(self):
        """退出应用"""
        if self.is_running:
            if messagebox.askyesno("确认", "测试正在执行中，确定要退出吗？"):
                self.is_running = False
                if self.screenshot_event:
                    self.screenshot_event.set()
                self.root.quit()
        else:
            self.root.quit()
    
    def run(self):
        """运行应用"""
        self.root.mainloop()

class AnnotationWindow:
    def __init__(self, parent, image_path, callback):
        self.parent = parent
        self.image_path = image_path
        self.callback = callback
        self.window = None
        self.canvas = None
        self.image = None
        self.tk_image = None
        self.drawing = False
        self.last_x = 0
        self.last_y = 0
        self.annotations = []  # 存储标注信息
        self.current_annotation = None
        
    def show(self):
        """显示标注窗口"""
        # 加载图像
        self.image = Image.open(self.image_path)
        self.draw = ImageDraw.Draw(self.image)
        
        # 创建窗口
        self.window = tk.Toplevel(self.parent)
        self.window.title("图像标注")
        self.window.state('zoomed')  # 最大化窗口
        
        # 创建工具栏
        toolbar = tk.Frame(self.window)
        toolbar.pack(side=tk.TOP, fill=tk.X)
        
        tk.Button(toolbar, text="圆形标注", command=self.start_circle_tool).pack(side=tk.LEFT, padx=5)
        tk.Button(toolbar, text="文字标注", command=self.start_text_tool).pack(side=tk.LEFT, padx=5)
        tk.Button(toolbar, text="撤销", command=self.undo_annotation).pack(side=tk.LEFT, padx=5)
        tk.Button(toolbar, text="保存", command=self.save_annotation, bg="lightgreen").pack(side=tk.RIGHT, padx=5)
        tk.Button(toolbar, text="取消", command=self.cancel_annotation).pack(side=tk.RIGHT, padx=5)
        
        # 创建画布
        self.canvas = tk.Canvas(self.window, bg='gray')
        self.canvas.pack(fill=tk.BOTH, expand=True)
        
        # 显示图像
        self.display_image()
        
        # 绑定鼠标事件
        self.canvas.bind("<Button-1>", self.on_mouse_down)
        self.canvas.bind("<B1-Motion>", self.on_mouse_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_mouse_up)
        self.canvas.bind("<Button-3>", self.on_right_click)  # 右键添加文字
        
        # 设置窗口为模态
        self.window.transient(self.parent)
        self.window.grab_set()
        
        # 居中显示
        self.window.update_idletasks()
        x = (self.window.winfo_screenwidth() // 2) - (self.window.winfo_width() // 2)
        y = (self.window.winfo_screenheight() // 2) - (self.window.winfo_height() // 2)
        self.window.geometry(f"+{x}+{y}")
        
    def display_image(self):
        """显示图像"""
        # 获取画布尺寸
        self.canvas.update()
        canvas_width = self.canvas.winfo_width()
        canvas_height = self.canvas.winfo_height()
        
        # 缩放图像以适应画布
        img_width, img_height = self.image.size
        scale = min(canvas_width/img_width, canvas_height/img_height)
        new_width = int(img_width * scale)
        new_height = int(img_height * scale)
        
        resized_image = self.image.resize((new_width, new_height), Image.LANCZOS)
        self.tk_image = ImageTk.PhotoImage(resized_image)
        
        # 在画布中心显示图像
        self.canvas.delete("image")
        x = (canvas_width - new_width) // 2
        y = (canvas_height - new_height) // 2
        self.canvas.create_image(x, y, anchor=tk.NW, image=self.tk_image, tags="image")
        
        self.scale = scale
        self.image_offset_x = x
        self.image_offset_y = y
    
    def start_circle_tool(self):
        """开始圆形标注工具"""
        self.current_tool = "circle"
        self.window.config(cursor="crosshair")
    
    def start_text_tool(self):
        """开始文字标注工具"""
        self.current_tool = "text"
        self.window.config(cursor="xterm")
    
    def undo_annotation(self):
        """撤销上一个标注"""
        if self.annotations:
            self.annotations.pop()
            self.redraw_image()
    
    def redraw_image(self):
        """重新绘制图像和所有标注"""
        # 重新加载原始图像
        self.image = Image.open(self.image_path)
        self.draw = ImageDraw.Draw(self.image)
        
        # 重新应用所有标注
        for annotation in self.annotations:
            if annotation['type'] == 'circle':
                self.draw.ellipse(annotation['coords'], outline=annotation['color'], width=3)
            elif annotation['type'] == 'text':
                self.draw.text(annotation['position'], annotation['text'], fill=annotation['color'])
        
        # 更新显示
        self.display_image()
    
    def on_mouse_down(self, event):
        """鼠标按下事件"""
        if not hasattr(self, 'current_tool'):
            return
            
        self.drawing = True
        self.last_x, self.last_y = event.x, event.y
        
        if self.current_tool == "circle":
            # 开始绘制圆形
            self.start_x, self.start_y = event.x, event.y
    
    def on_mouse_drag(self, event):
        """鼠标拖拽事件"""
        if not self.drawing or not hasattr(self, 'current_tool') or self.current_tool != "circle":
            return
            
        # 重新绘制图像以清除之前的预览
        self.redraw_image()
        
        # 绘制当前圆形预览
        x0, y0 = self.start_x, self.start_y
        x1, y1 = event.x, event.y
        
        # 转换为图像坐标
        img_x0 = (x0 - self.image_offset_x) / self.scale
        img_y0 = (y0 - self.image_offset_y) / self.scale
        img_x1 = (x1 - self.image_offset_x) / self.scale
        img_y1 = (y1 - self.image_offset_y) / self.scale
        
        # 在原图上绘制预览
        preview_draw = ImageDraw.Draw(self.image)
        preview_draw.ellipse([img_x0, img_y0, img_x1, img_y1], outline="red", width=3)
        self.display_image()
    
    def on_mouse_up(self, event):
        """鼠标释放事件"""
        if not self.drawing or not hasattr(self, 'current_tool') or self.current_tool != "circle":
            return
            
        self.drawing = False
        
        # 保存圆形标注
        x0, y0 = self.start_x, self.start_y
        x1, y1 = event.x, event.y
        
        # 转换为图像坐标
        img_x0 = (x0 - self.image_offset_x) / self.scale
        img_y0 = (y0 - self.image_offset_y) / self.scale
        img_x1 = (x1 - self.image_offset_x) / self.scale
        img_y1 = (y1 - self.image_offset_y) / self.scale
        
        # 保存标注信息
        self.annotations.append({
            'type': 'circle',
            'coords': [img_x0, img_y0, img_x1, img_y1],
            'color': 'red'
        })
        
        # 重新绘制以确保标注正确显示
        self.redraw_image()
    
    def on_right_click(self, event):
        """右键点击事件 - 添加文字"""
        if not hasattr(self, 'current_tool') or self.current_tool != "text":
            return
        
        # 弹出输入对话框
        from tkinter.simpledialog import askstring
        text = askstring("添加文字", "请输入标注文字:", parent=self.window)
        if text:
            # 转换为图像坐标
            img_x = (event.x - self.image_offset_x) / self.scale
            img_y = (event.y - self.image_offset_y) / self.scale
            
            # 保存文字标注信息
            self.annotations.append({
                'type': 'text',
                'position': (img_x, img_y),
                'text': text,
                'color': 'red'
            })
            
            # 重新绘制
            self.redraw_image()
    
    def save_annotation(self):
        """保存标注"""
        try:
            # 保存标注后的图像
            output_path = f"annotated_{int(time.time())}.png"
            self.image.save(output_path)
            
            # 关闭窗口
            self.window.destroy()
            
            # 调用回调函数
            if self.callback:
                self.callback(output_path)
                
        except Exception as e:
            messagebox.showerror("错误", f"保存标注失败: {str(e)}")
    
    def cancel_annotation(self):
        """取消标注"""
        if messagebox.askyesno("确认", "确定要取消标注吗？未保存的更改将丢失。"):
            self.window.destroy()

if __name__ == "__main__":
    app = AutomationTestTool()
    app.run()