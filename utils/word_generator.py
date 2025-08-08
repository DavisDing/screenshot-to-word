# utils/word_generator.py

import os
from docx import Document
from docx.shared import Inches
from tkinter import messagebox
from utils.path_utils import get_base_path

class WordGenerator:
    def __init__(self, logger, root):
        self.logger = logger
        self.root = root
        base_dir = get_base_path()
        self.output_dir = os.path.join(base_dir, "word_output")
        os.makedirs(self.output_dir, exist_ok=True)

    def add_image_to_word(self, case_name, checkpoint, image_path, step_note=""):
        """
        新建或追加Word文档，插入用例名、验证点、（可选）步骤说明和图片
        """

        doc_path = os.path.join(self.output_dir, f"{case_name}.docx")

        if os.path.exists(doc_path):
            try:
                doc = Document(doc_path)
                self.logger.log(f"打开已有Word文件：{doc_path}")
            except Exception as e:
                self._show_error(f"打开Word文件失败：{e}")
                return
        else:
            doc = Document()
            doc.add_heading(case_name, level=1)
            doc.add_paragraph(checkpoint)

        if step_note:
            doc.add_paragraph(step_note)

        # 插入图片，自动适应宽度
        try:
            doc.add_picture(image_path, width=Inches(5.5))
            doc.add_paragraph()  # 空行
            doc.save(doc_path)
            self.logger.log(f"插入截图并保存Word：{doc_path}")
        except Exception as e:
            self._show_error(f"写入Word文件失败：{e}")

    def _show_error(self, msg):
        def _task():
            if self.root.winfo_exists():
                self.root.attributes('-topmost', True)
                messagebox.showerror("错误", msg, parent=self.root)
                self.root.attributes('-topmost', False)
        if self.root:
            self.root.after(0, _task)